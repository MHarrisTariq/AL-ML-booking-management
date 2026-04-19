"""
Build SwyftBooking_Codebook.docx: explanations + full source listings.
Requires: pip install python-docx
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "SwyftBooking_Codebook.docx"

# (relative path, short explanation)
SECTIONS: list[tuple[str, str]] = [
    (
        "README.md",
        "Overview of SwyftBooking: hybrid channel manager, quick start (memory mode vs Mongo/Redis), and repo layout.",
    ),
    (
        "docker-compose.yml",
        "Full stack services: MongoDB, Redis, Zookeeper/Kafka, FastAPI AI, NestJS backend, React frontend; networking and env vars.",
    ),
    (
        "backend/package.json",
        "Backend npm scripts (including start:dev:memory), dependencies, and Jest config.",
    ),
    (
        "backend/Dockerfile",
        "Multi-stage Node image for production backend build.",
    ),
    (
        "backend/src/main.ts",
        "NestJS bootstrap: global prefix /api, CORS, validation pipe, listen port.",
    ),
    (
        "backend/src/app.module.ts",
        "Root module: Config, conditional Mongoose vs in-memory booking, Redis, Kafka, Booking vs BookingMemory.",
    ),
    (
        "backend/src/app.controller.ts",
        "Root and /health endpoints for API liveness.",
    ),
    (
        "backend/src/app.service.ts",
        "Health payload for monitoring.",
    ),
    (
        "backend/src/app.controller.spec.ts",
        "Unit tests for AppController.",
    ),
    (
        "backend/src/config/configuration.ts",
        "Centralized env: port, Mongo URI, Redis, Kafka brokers, AI base URL.",
    ),
    (
        "backend/src/kafka/topics.ts",
        "Kafka topic name constants (booking.*, sync.*, risk.evaluated).",
    ),
    (
        "backend/src/kafka/kafka.service.ts",
        "KafkaJS producer: connect, send JSON payloads (tolerates broker absence in dev).",
    ),
    (
        "backend/src/kafka/kafka.module.ts",
        "Global Kafka module exporting KafkaService.",
    ),
    (
        "backend/src/redis/redis.service.ts",
        "Redis client: real ioredis, or ioredis-mock when SWYFT_DEV_MEMORY=1.",
    ),
    (
        "backend/src/redis/redis.module.ts",
        "Global Redis module.",
    ),
    (
        "backend/src/lock/lock.service.ts",
        "Distributed lock pattern: SET NX EX + Lua release for listing/date key.",
    ),
    (
        "backend/src/lock/lock.module.ts",
        "Lock module exporting LockService.",
    ),
    (
        "backend/src/ai/ai-client.service.ts",
        "HTTP client to FastAI services: risk, availability probability, demand, sync interval.",
    ),
    (
        "backend/src/ai/ai.module.ts",
        "HttpModule + AiClientService.",
    ),
    (
        "backend/src/decision/decision-engine.service.ts",
        "Hybrid decision: demand forecast, availability probability threshold, risk score actions (approve/block/delay).",
    ),
    (
        "backend/src/decision/decision.module.ts",
        "DecisionEngineService wired with AiModule.",
    ),
    (
        "backend/src/availability/availability.service.ts",
        "Mongo-backed overlap check for confirmed/pending bookings.",
    ),
    (
        "backend/src/availability/availability.module.ts",
        "Mongoose Booking model registration for availability.",
    ),
    (
        "backend/src/availability/availability-memory.service.ts",
        "In-memory overlap check using dev booking store.",
    ),
    (
        "backend/src/availability/availability-memory.module.ts",
        "Exports AvailabilityMemoryService.",
    ),
    (
        "backend/src/booking/schemas/booking.schema.ts",
        "Mongoose schema: listing, guest, dates, idempotency, status, price.",
    ),
    (
        "backend/src/booking/dto/create-booking.dto.ts",
        "Validation DTO for POST /bookings.",
    ),
    (
        "backend/src/booking/booking.controller.ts",
        "REST: create booking, list by listingId, get by id (route order fixed for listing path).",
    ),
    (
        "backend/src/booking/booking.service.ts",
        "Deterministic flow: idempotency (Redis), AI decision, lock, Mongo conflict check, persist, Kafka, sync plan.",
    ),
    (
        "backend/src/booking/booking.module.ts",
        "Mongoose booking feature module (production path).",
    ),
    (
        "backend/src/dev/dev-booking.store.ts",
        "In-process booking rows when SWYFT_DEV_MEMORY=1.",
    ),
    (
        "backend/src/booking/booking-memory.service.ts",
        "Same booking logic as booking.service but uses dev store + AvailabilityMemoryService.",
    ),
    (
        "backend/src/booking/booking-memory.module.ts",
        "Registers BookingMemoryService as BookingService token for controller injection.",
    ),
    (
        "backend/src/sync/sync-orchestrator.service.ts",
        "Logs and emits sync.requested to Kafka with AI-suggested interval metadata.",
    ),
    (
        "backend/src/sync/sync.module.ts",
        "Sync orchestrator module.",
    ),
    (
        "frontend/package.json",
        "Frontend scripts and Vite/React dependencies.",
    ),
    (
        "frontend/vite.config.ts",
        "Vite config and dev proxy /api to backend port 3000.",
    ),
    (
        "frontend/index.html",
        "SPA entry HTML.",
    ),
    (
        "frontend/src/main.tsx",
        "React root mount.",
    ),
    (
        "frontend/src/App.tsx",
        "Host dashboard: health, create booking, list by listing, refresh; calls /api.",
    ),
    (
        "frontend/src/App.css",
        "Layout and styling for dashboard cards.",
    ),
    (
        "frontend/src/index.css",
        "Global reset/body.",
    ),
    (
        "frontend/src/vite-env.d.ts",
        "Vite client types.",
    ),
    (
        "frontend/Dockerfile",
        "Static build + nginx proxy /api to backend service.",
    ),
    (
        "frontend/nginx.conf",
        "Nginx SPA + API reverse proxy for containerized frontend.",
    ),
    (
        "ai-services/main.py",
        "FastAPI: /risk-score (LightGBM or heuristic), /availability-probability, /sync-interval, /demand-forecast, /platform-reliability.",
    ),
    (
        "ai-services/requirements.txt",
        "Python dependencies for AI service.",
    ),
    (
        "ai-services/Dockerfile",
        "Python 3.11 slim image running uvicorn.",
    ),
    (
        "ml/train_risk_model.py",
        "Synthetic dataset, LightGBM/sklearn training, joblib export to ai-services/models, optional MLflow.",
    ),
    (
        "ml/pipelines/airflow_train_dag.py",
        "Airflow DAG stub calling training script (requires Airflow in target env).",
    ),
    (
        "streaming/kafka_client.py",
        "Python Kafka producer helper and topic constants aligned with backend.",
    ),
    (
        "streaming/requirements.txt",
        "kafka-python for streaming helpers.",
    ),
    (
        "data/schemas/booking_record.json",
        "JSON Schema sample for booking records.",
    ),
    (
        "infra/k8s/backend-deployment.yaml",
        "Example Kubernetes Deployment + Service for backend.",
    ),
    (
        "infra/k8s/ai-services-deployment.yaml",
        "Example Deployment + Service for AI microservice.",
    ),
    (
        "infra/k8s/frontend-deployment.yaml",
        "Example Deployment + Service for static frontend.",
    ),
    (
        ".gitignore",
        "Ignore patterns for node_modules, dist, env files, Python cache.",
    ),
]


def set_code_font(run, size_pt: float = 8) -> None:
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_code_block(doc: Document, text: str) -> None:
    if not text.endswith("\n"):
        text += "\n"
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_code_font(run, 8)


def main() -> None:
    doc = Document()
    title = doc.add_heading("SwyftBooking — Source Codebook", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph(
        "This document lists the main project files with a short explanation of each, "
        "followed by the file contents as of generation time. "
        "SwyftBooking is a hybrid channel-manager starter: NestJS deterministic core, "
        "FastAPI AI services, React UI, optional Kafka/Redis/Mongo via Docker."
    )
    sub.paragraph_format.space_after = Pt(12)

    missing: list[str] = []
    for rel, explain in SECTIONS:
        path = ROOT / rel
        doc.add_page_break()
        doc.add_heading(rel.replace("/", " → "), level=1)
        p = doc.add_paragraph(explain)
        p.paragraph_format.space_after = Pt(6)
        if not path.is_file():
            missing.append(rel)
            doc.add_paragraph(f"[File not found on disk: {rel}]")
            continue
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            doc.add_paragraph(f"[Could not read file: {e}]")
            continue
        add_code_block(doc, raw)

    if missing:
        doc.add_page_break()
        doc.add_heading("Missing files", level=1)
        doc.add_paragraph(
            "These paths were not found (optional or renamed): " + ", ".join(missing)
        )

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
