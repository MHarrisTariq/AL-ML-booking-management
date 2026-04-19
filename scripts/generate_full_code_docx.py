from __future__ import annotations

import os
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "SwyftBooking_All_Code_and_Explanation.docx"

EXCLUDE_DIRS = {
    ".git",
    "node_modules",
    "dist",
    "build",
    "coverage",
    "__pycache__",
    ".venv",
    "venv",
    ".mypy_cache",
    ".pytest_cache",
    ".next",
    ".vite",
    "analysis1_docx_unzipped",
}

EXCLUDE_FILES = {
    "Analysis 1 (1).docx",
    "Analysis 1 (1).zip",
    "Analysis 1 (1).extracted.txt",
    "AI ML Booking Managment 1.docx",
    "SwyftBooking_Codebook.docx",
    ".DS_Store",
    str(OUT.name),
}

MAX_CHARS_PER_FILE = 250_000


def iter_text_files(root: Path):
    for dirpath, dirnames, filenames in os.walk(root):
        dp = Path(dirpath)
        dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIRS]
        for fn in filenames:
            p = dp / fn
            if p.name in EXCLUDE_FILES:
                continue
            if p.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".pdf", ".docx", ".zip"}:
                continue
            yield p


def read_text(p: Path) -> str:
    try:
        data = p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        data = p.read_text(encoding="utf-8", errors="replace")
    if len(data) > MAX_CHARS_PER_FILE:
        data = data[:MAX_CHARS_PER_FILE] + "\n\n... [TRUNCATED: file too large for DOCX output] ...\n"
    return sanitize_xml_text(data)


def sanitize_xml_text(s: str) -> str:
    # WordprocessingML rejects NULL bytes and most control chars.
    # Keep tab/newline/carriage return; replace others with U+FFFD.
    out_chars: list[str] = []
    for ch in s:
        o = ord(ch)
        if ch in ("\t", "\n", "\r"):
            out_chars.append(ch)
        elif o == 0x00:
            out_chars.append("\uFFFD")
        elif 0x01 <= o <= 0x08:
            out_chars.append("\uFFFD")
        elif 0x0B <= o <= 0x0C:
            out_chars.append("\uFFFD")
        elif 0x0E <= o <= 0x1F:
            out_chars.append("\uFFFD")
        else:
            out_chars.append(ch)
    return "".join(out_chars)


def main() -> Path:
    doc = docx.Document()

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    if "CodeBlock" not in [s.name for s in styles]:
        code_style = styles.add_style("CodeBlock", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_style.font.size = Pt(9)

    def h(text: str, level: int):
        doc.add_heading(text, level=level)

    h("SwyftBooking — Full Code + Explanation", 0)
    doc.add_paragraph(
        "This document contains (a) an explanation of the applied hardening changes and "
        "(b) a full listing of the project code (excluding node_modules and other generated/binary artifacts)."
    )

    h("Explanation of Changes Applied", 1)
    sections = [
        (
            "AI resilience",
            [
                "Added AI request timeout + retry with deterministic fallbacks so the booking flow never fails if the AI service is slow/down.",
                "Decision flow now uses safe AI calls; sync interval call uses safe fallback as well.",
                "Files: backend/src/ai/ai-client.service.ts, backend/src/decision/decision-engine.service.ts, backend/src/booking/booking.service.ts",
            ],
        ),
        (
            "Database performance (availability conflict)",
            [
                "Replaced linear cursor scan with indexed overlap query.",
                "Added Mongo index on (listingId, startDate, endDate, status).",
                "Files: backend/src/availability/availability.service.ts, backend/src/booking/schemas/booking.schema.ts",
            ],
        ),
        (
            "Kafka: consumers + retry/DLQ",
            [
                "Added Kafka consumer service (event-driven behavior) and retry+DLQ topics booking.retry and booking.dlq.",
                "Added basic idempotency per topic/eventId using Redis keys.",
                "Files: backend/src/kafka/kafka-consumer.service.ts, backend/src/kafka/topics.ts, backend/src/kafka/kafka.module.ts",
            ],
        ),
        (
            "Health endpoints",
            [
                "Added /ai-health and /kafka-health endpoints.",
                "Files: backend/src/app.controller.ts, backend/src/kafka/kafka.service.ts",
            ],
        ),
        (
            "Locking hardening",
            [
                "Introduced Redlock-based locking (with a dev-memory fallback).",
                "Files: backend/src/lock/lock.service.ts, backend/package.json",
            ],
        ),
    ]
    for title, bullets in sections:
        doc.add_paragraph(title, style="List Bullet")
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet 2")

    h("Full Source Code Listing", 1)
    doc.add_paragraph(
        "Included: backend/, frontend/ (source), ai-services/, ml/, streaming/, infra/, scripts/, and root config files. "
        "Excluded: node_modules/, generated outputs, and binary files (docx/zip/images)."
    )

    files = sorted(iter_text_files(ROOT), key=lambda p: str(p).lower())
    for p in files:
        rel = p.relative_to(ROOT)
        # extra defense: skip node_modules even if present via symlinks
        if "node_modules" in rel.parts:
            continue
        # skip empty
        try:
            if p.stat().st_size == 0:
                continue
        except OSError:
            continue

        h(str(rel).replace("/", "\\"), 2)
        text = read_text(p)
        doc.add_paragraph(f"Path: {rel} | Characters: {len(text):,}")
        for line in text.splitlines():
            doc.add_paragraph(line, style="CodeBlock")

    if OUT.exists():
        OUT.unlink()
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    out = main()
    print(out)

